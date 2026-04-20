import streamlit as st
import google.generativeai as genai
from google.generativeai.types import HarmCategory, HarmBlockThreshold
import io, time, re
from PIL import Image
import fitz # PyMuPDF
from docx import Document
from docx.shared import Inches

# --- CẤU HÌNH HỆ THỐNG ---
st.set_page_config(page_title="Gemini Next-Gen Translator", layout="wide", page_icon="🚀")

# --- HÀM TỰ ĐỘNG CẬP NHẬT MODEL MỚI NHẤT ---
def get_latest_gemini_model():
    try:
        # Lấy danh sách tất cả model khả dụng
        available_models = []
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods and 'gemini' in m.name:
                available_models.append(m.name)
        
        if not available_models:
            return "models/gemini-1.5-pro" # Dự phòng tối thiểu

        # Sắp xếp theo thứ tự phiên bản (ví dụ: 3.1 > 2.0 > 1.5)
        # Tìm các con số trong tên model để so sánh
        def natural_sort_key(s):
            return [int(text) if text.isdigit() else text.lower() for text in re.split('([0-9]+)', s)]
        
        available_models.sort(key=natural_sort_key, reverse=True)
        
        # Ưu tiên bản "pro" nếu có nhiều biến thể cùng phiên bản
        for m in available_models:
            if "pro" in m.lower():
                return m
                
        return available_models[0]
    except Exception as e:
        st.warning(f"⚠️ Không thể quét danh sách model: {e}. Dùng mặc định.")
        return "models/gemini-1.5-pro"

# --- KẾT NỐI API ---
if "GEMINI_API_KEY" not in st.secrets:
    st.error("❌ Thiếu GEMINI_API_KEY trong Streamlit Secrets!")
    st.stop()

genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
LATEST_MODEL = get_latest_gemini_model()

st.sidebar.success(f"🤖 Đang sử dụng phiên bản: {LATEST_MODEL}")

safety_settings = {
    HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
}

# --- ENGINE DỊCH THUẬT THÔNG MINH ---

def translate_logic(prompt_data):
    """Gửi yêu cầu dịch và tự động xử lý các lỗi nghẽn mạch"""
    model = genai.GenerativeModel(LATEST_MODEL)
    for attempt in range(5):
        try:
            res = model.generate_content(prompt_data, safety_settings=safety_settings)
            if res and res.text:
                return res.text
        except Exception as e:
            err_str = str(e)
            if "429" in err_str or "quota" in err_str.lower():
                wait = 20 + (attempt * 10)
                st.toast(f"⏳ Hạn mức API tạm hết, đợi {wait}s...", icon="⏳")
                time.sleep(wait)
            else:
                st.error(f"❌ Lỗi: {err_str}")
                time.sleep(5)
    return "[Lỗi: Không thể hoàn thành đoạn dịch này]"

def process_heavy_text(text, instruction):
    """Chia nhỏ 30.000 từ thành các block để dịch không bị mất chữ"""
    # Ngay cả Gemini 3.1 cũng có giới hạn Token đầu ra, nên chia block 12k ký tự là an toàn nhất
    chunks = [text[i:i+12000] for i in range(0, len(text), 12000)]
    translated_parts = []
    
    for c in chunks:
        p = f"{instruction}\n\nNỘI DUNG GỐC:\n{c}"
        part = translate_logic(p)
        translated_parts.append(part)
        time.sleep(1) # Nghỉ ngắn tránh lỗi spam
    
    return "\n".join(translated_parts)

# --- XỬ LÝ FILE ---

def extract_pdf(f):
    doc = fitz.open(stream=f.read(), filetype="pdf")
    results = []
    for page in doc:
        # Lấy văn bản
        text = page.get_text().strip()
        if text: results.append({"type": "text", "val": text})
        # Lấy hình ảnh
        for img_info in page.get_images():
            try:
                xref = img_info[0]
                img_data = doc.extract_image(xref)
                if img_data["size"] > 10000:
                    pil_img = Image.open(io.BytesIO(img_data["image"]))
                    results.append({"type": "image", "val": pil_img})
            except: pass
    return results

# --- GIAO DIỆN NGƯỜI DÙNG ---
st.title("🛡️ Siêu Trợ Lý Gemini 3.1 Auto-Update")
st.subheader("Dịch thuật công nghiệp không giới hạn số từ")

with st.sidebar:
    st.header("Cấu hình dịch")
    instr = st.text_area("Chỉ thị dịch:", "Dịch sang tiếng Việt mượt mà, văn phong chuyên nghiệp, giữ nguyên ý nghĩa gốc.")
    glossary = st.text_area("Thuật ngữ:", "VD: Blockchain -> Chuỗi khối")

up_files = st.file_uploader("Tải lên tài liệu (PDF/Docx/TXT):", accept_multiple_files=True)

if st.button("🚀 BẮT ĐẦU DỊCH NGAY") and up_files:
    for f in up_files:
        st.write(f"📂 **Đang xử lý file:** {f.name}")
        content_list = extract_pdf(f)
        
        translated_content = []
        p_bar = st.progress(0)
        
        for idx, item in enumerate(content_list):
            if item['type'] == 'text':
                res = process_heavy_text(item['val'], f"{instr}\nThuật ngữ: {glossary}")
                translated_content.append({"type": "text", "val": res})
            elif item['type'] == 'image':
                # Dịch ảnh dùng Multimodal
                res = translate_logic([f"Dịch các chữ trong ảnh này sang tiếng Việt. {instr}", item['val']])
                translated_content.append({"type": "image", "val": item['val'], "trans": res})
            
            p_bar.progress((idx + 1) / len(content_list))

        # Xuất file Word
        out_doc = Document()
        for c in translated_content:
            if c['type'] == 'text':
                out_doc.add_paragraph(c['val'])
            else:
                img_stream = io.BytesIO()
                c['val'].save(img_stream, format='PNG')
                out_doc.add_picture(img_stream, width=Inches(5))
                out_doc.add_paragraph(f"[Bản dịch ảnh]: {c['trans']}").italic = True
        
        final_bio = io.BytesIO()
        out_doc.save(final_bio)
        st.download_button(f"⬇️ Tải bản dịch {f.name}", final_bio.getvalue(), f"Translated_{f.name}.docx")
        st.success(f"✅ Đã xong: {f.name}")
